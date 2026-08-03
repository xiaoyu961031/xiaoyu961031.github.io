from __future__ import annotations

import csv
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT.parent / "output" / "cv"
CSV_PATH = ROOT / "public" / "publications.csv"

FONT = "Arial"
INK = RGBColor(24, 31, 29)
MUTED = RGBColor(88, 96, 92)
ACCENT = RGBColor(22, 124, 88)
LINE = "D7DDDA"


def set_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return run


def set_cellless_rule(paragraph, color=LINE, size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_hyperlink(paragraph, text, url, color=ACCENT, underline=False, size=9):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    fonts.set(qn("w:cs"), FONT)
    r_pr.append(fonts)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), str(color))
    r_pr.append(color_el)
    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(size_el)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    fonts.set(qn("w:cs"), FONT)
    r_pr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "16")
    r_pr.append(size)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "585F5C")
    r_pr.append(color)
    run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def add_numbering(doc, fmt="decimal", text="%1.", left=360, hanging=260):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abs or [0]) + 1
    num_id = max(existing_num or [0]) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "space")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.append(ind)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)


def configure_document(doc, margin=0.68, body_size=9.3):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.30)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal.font.size = Pt(body_size)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.06
    for style_name, size, before, after in [("Heading 1", 12, 9, 4), ("Heading 2", 10.5, 7, 3), ("Heading 3", 9.5, 5, 2)]:
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    set_font(footer.add_run("Xiaoyu Wu  |  "), 8, color=MUTED)
    add_page_field(footer)


def add_header(doc, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run("Xiaoyu Wu"), 22, bold=True, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(subtitle), 10.5, bold=True, color=ACCENT)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    set_font(p.add_run("Singapore  |  +65 8067 0533  |  "), 8.8, color=MUTED)
    add_hyperlink(p, "xiaoyuwu742@gmail.com", "mailto:xiaoyuwu742@gmail.com", size=8.8)
    set_font(p.add_run("  |  "), 8.8, color=MUTED)
    add_hyperlink(p, "LinkedIn", "https://www.linkedin.com/in/xiaoyu-wu-8538b7166/", size=8.8)
    set_font(p.add_run("  |  "), 8.8, color=MUTED)
    add_hyperlink(p, "GitHub", "https://github.com/xiaoyu961031", size=8.8)
    set_font(p.add_run("  |  "), 8.8, color=MUTED)
    add_hyperlink(p, "ORCID", "https://orcid.org/0009-0004-5625-1180", size=8.8)
    set_cellless_rule(p)


def heading(doc, text, level=1):
    sizes = {1: 12, 2: 10.5, 3: 9.5}
    before = {1: 9, 2: 7, 3: 5}
    after = {1: 4, 2: 3, 3: 2}
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before[level])
    paragraph.paragraph_format.space_after = Pt(after[level])
    paragraph.paragraph_format.keep_with_next = True
    set_font(paragraph.add_run(text), sizes[level], bold=True, color=ACCENT)
    return paragraph


def body(doc, text, size=None, italic=False, color=None, after=None):
    p = doc.add_paragraph()
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    set_font(p.add_run(text), size, italic=italic, color=color)
    return p


def bullet(doc, text, num_id, size=None, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    p.paragraph_format.space_after = Pt(2.5)
    set_font(p.add_run("•  "), size)
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), size, bold=True)
        set_font(p.add_run(text[len(bold_prefix):]), size)
    else:
        set_font(p.add_run(text), size)
    return p


def role_entry(doc, dates, title, organization, details=None, bullets=None, bullet_id=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(title), 9.4, bold=True)
    set_font(p.add_run(f"  |  {organization}"), 9.4, color=MUTED)
    date_run = p.add_run(f"  |  {dates}")
    set_font(date_run, 8.8, bold=True, color=ACCENT)
    if details:
        body(doc, details, 8.8, color=MUTED, after=2)
    for item in bullets or []:
        bullet(doc, item, bullet_id, 8.9)


def read_publications():
    with CSV_PATH.open(encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def clean_authors(authors):
    return authors.replace("[cf]", "†").replace("[cor]", "*").replace("[cf,cor]", "†*")


def add_publication(doc, paper, index):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.24)
    p.paragraph_format.first_line_indent = Inches(-0.21)
    p.paragraph_format.space_after = Pt(3.4)
    p.paragraph_format.line_spacing = 1.02
    set_font(p.add_run(f"{index}. "), 8.15)
    authors = clean_authors(paper["authors"])
    chunks = re.split(r"(Xiaoyu Wu(?:†\*|†|\*)?)", authors)
    for chunk in chunks:
        if not chunk:
            continue
        set_font(p.add_run(chunk), 8.15, bold=chunk.startswith("Xiaoyu Wu"))
    set_font(p.add_run(f'. “{paper["title"]}.” '), 8.15)
    set_font(p.add_run(paper["journal"]), 8.15, italic=True)
    status = f', {paper["year"]}' + (f' ({paper["status_or_note"]})' if paper["status_or_note"] else "")
    set_font(p.add_run(status), 8.15, color=MUTED)
    if paper["doi_or_link"]:
        set_font(p.add_run("  "), 8.15)
        add_hyperlink(p, "DOI", paper["doi_or_link"], size=8.15)


def add_in_progress(doc):
    heading(doc, "Manuscripts Under Review", 1)
    entries = [
        ("Linker Functionalization in HKUST-1 Metal-Organic Frameworks: Insights into Synthesizability and Hydrolytic Properties from Machine Learning Potentials", "Rui Zheng; Xiaoyu Wu*; Jianwen Jiang*", "Journal of the American Chemical Society"),
        ("Spatially Aware Deep Learning for Isotopologue Separation: Predicting 12CH4/13CH4 Selectivity in Metal-Organic Frameworks via Grid-Crystal Residual Networks", "Chenrui Li†; Xiaoyu Wu*; Tianyu Guo; Linjiang Chen; Lin-Bing Sun*; Zhenghao Wu*; Lifeng Ding*", "Separation and Purification Technology"),
    ]
    for title, authors, journal in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        set_font(p.add_run(title + ". "), 8.7, bold=True)
        set_font(p.add_run(authors + ". "), 8.5)
        set_font(p.add_run(f"Under review at {journal}."), 8.5, italic=True, color=MUTED)


def build_academic(publications):
    doc = Document()
    configure_document(doc, margin=0.64, body_size=9.1)
    bullet_id = add_numbering(doc, "bullet", "•", 300, 180)
    pub_num_id = add_numbering(doc, "decimal", "%1.", 340, 260)
    add_header(doc, "PhD, MRSC  |  Research Fellow  |  Computational Chemistry and AI for Reticular Chemistry")
    heading(doc, "Academic Profile")
    body(doc, "Computational chemist developing molecular simulation, high-throughput screening, and machine-learning approaches for the digital discovery of nanoporous materials. Research spans metal-organic and covalent organic frameworks, porous liquids, adsorption and separation, framework stability and flexibility, and chemically informed AI. Experienced in interdisciplinary collaboration, student mentoring, scientific communication, and large-scale computational research.", 9.1)
    heading(doc, "Appointments")
    role_entry(doc, "2023-present", "Research Fellow", "National University of Singapore", "Department of Chemical and Biomolecular Engineering; supervised by Prof. Jianwen Jiang.")
    heading(doc, "Education")
    role_entry(doc, "2023", "PhD in Chemistry", "University of Liverpool, United Kingdom", "Supervised by Prof. Andrew I. Cooper and Prof. Lifeng Ding.")
    role_entry(doc, "2018", "BSc in Chemistry", "Nanjing Tech University, China")
    heading(doc, "Research Focus")
    for item in [
        "AI for reticular chemistry: forward and inverse learning, active learning, deep learning, large language models, and machine-learning potentials.",
        "Molecular and multiscale simulation: adsorption, separation, transport, flexibility, stability, and structure-property relationships in porous materials.",
        "Sustainable applications: carbon capture, gas upgrading, hydrocarbon and isotope separations, sensing, and radionuclide sequestration.",
    ]:
        bullet(doc, item, bullet_id, 8.9)
    add_in_progress(doc)
    heading(doc, f"Peer-Reviewed Publications ({len(publications)})")
    for index, paper in enumerate(publications, 1):
        add_publication(doc, paper, index)
    heading(doc, "Teaching and Mentorship")
    bullet(doc, "Teaching assistant and co-supervisor across Introduction to Chemistry; Introductory Inorganic, Physical, and Organic Chemistry I-II; Laboratory Techniques; Preparative Chemistry; and Analytical Chemistry.", bullet_id, 8.8)
    bullet(doc, "Mentored MSc, BSc, and PhD researchers on MOF/COF simulation, adsorption and separation, framework mechanics and stability, machine learning, and LLM-assisted literature mining.", bullet_id, 8.8)
    bullet(doc, "Mentees include Rui Zheng (current PhD student, NUS), Tianyu Guo (now PhD researcher, HKUST), Mingyue Ma, Yixu Cheng, and Daeun Jung.", bullet_id, 8.8)
    heading(doc, "Selected Conference Presentations")
    conferences = [
        "AIChE Annual Meeting, Boston, United States (2025) - oral: Rejuvenating digital discoveries of MOFs through screening, featurization, and diversity-driven machine learning.",
        "AIChE Annual Meeting, Boston, United States (2025) - oral: Guest-induced structural transition in a flexible MOF for C4 alkane/alkene separation.",
        "AIChE Annual Meeting, Boston, United States (2025) - poster: Large language models for MOF hydrophobicity prediction.",
        "MOF 2024, Singapore (2024) - poster: Construction of hypothetical MOFs driven by reticular design.",
        "AIChE Annual Meeting, San Diego, United States (2024) - poster: Cross-database discovery of open-Cu(II) MOFs for biogas upgrading.",
    ]
    for item in conferences:
        bullet(doc, item, bullet_id, 8.8)
    heading(doc, "Research Funding and Computational Resources")
    for item in [
        "Participant, LCER Phase 2 programme (SGD 23,481,120).",
        "Participant, LCER Phase 1 projects LCERFI01-0015 U2102d2004 and LCERFI01-0033 U2102d2006 (SGD 7,534,320).",
        "Participant, NRF-CRP26-2021RS-0002 (SGD 8,871,670).",
        "Awarded large-scale NSCC and NUS CRP computational allocations, including multi-million CPU-hour and substantial GPU-hour projects.",
    ]:
        bullet(doc, item, bullet_id, 8.8)
    heading(doc, "Peer Review")
    body(doc, "Reviewer for Chemical Engineering Journal, Water Research, Journal of Colloid and Interface Science, Computational Materials Science, Journal of Chemical Information and Modeling, Journal of Physical Chemistry B/C, ACS Applied Materials & Interfaces, Energy Conversion and Management, and related journals.", 8.8)
    heading(doc, "Selected Research Outputs")
    for item in [
        "Precision-Engineered MOF database and reverse-topological design workflow (94,823 structures).",
        "Open-copper-site MOF databases and cross-diversity machine-learning models.",
        "MD-informed deep-learning dataset for synthesizable MOF discovery.",
        "Active-learning dataset for precombustion CO2 capture; Pore+ descriptors; MOF-KAN; fine-tuned LLMs for MOF hydrophobicity.",
    ]:
        bullet(doc, item, bullet_id, 8.8)
    heading(doc, "Professional Membership")
    body(doc, "Member of the Royal Society of Chemistry (MRSC).", 8.8)
    return doc


def build_industry(publications):
    doc = Document()
    configure_document(doc, margin=0.68, body_size=9.5)
    bullet_id = add_numbering(doc, "bullet", "•", 320, 190)
    add_header(doc, "Computational Chemist  |  AI and Molecular Simulation for Materials Discovery")
    heading(doc, "Professional Summary")
    body(doc, "Computational chemist with a PhD and a track record of translating molecular simulation and machine learning into practical materials-discovery workflows. Experienced in high-throughput screening, scientific data analysis, active and deep learning, LLM adaptation, model interpretation, and cross-disciplinary delivery. Author of 29 peer-reviewed papers, including 18 co-first, first, or corresponding-authored works.", 9.4)
    heading(doc, "Core Capabilities")
    for item in [
        "Materials discovery: molecular dynamics, Monte Carlo simulation, adsorption and separation modelling, stability and flexibility analysis, high-throughput screening.",
        "AI and data: machine learning, deep learning, active learning, large language models, machine-learning potentials, feature engineering, model transferability and interpretation.",
        "Delivery: scientific programming, reproducible workflows, large-scale CPU/GPU computing, technical writing, publication leadership, mentoring, and interdisciplinary collaboration.",
    ]:
        bullet(doc, item, bullet_id, 9.1)
    heading(doc, "Professional Experience")
    role_entry(doc, "2023-present", "Research Fellow", "National University of Singapore", bullets=[
        "Develop end-to-end digital discovery workflows linking chemical questions, simulation data, predictive models, and experimentally relevant candidates.",
        "Built and analysed materials datasets at scales up to 94,823 designed MOFs; developed transferable and stability-aware models across heterogeneous chemical spaces.",
        "Combined molecular dynamics with high-fidelity deep learning to assess MOF synthesizability; work selected as an Advanced Functional Materials cover article.",
        "Led or co-led projects in carbon capture, gas upgrading, molecular and isotope separation, porous-material stability, flexibility, sensing, and photocatalysis.",
        "Mentor PhD, MSc, and BSc researchers; collaborate with simulation, synthesis, membrane, and data-science teams across institutions.",
    ], bullet_id=bullet_id)
    doc.add_page_break()
    heading(doc, "Selected Project Outcomes")
    projects = [
        ("Synthesizability-aware discovery", "Integrated molecular-dynamics evidence with deep learning to prioritise chemically realistic MOFs.", "https://doi.org/10.1002/adfm.202519565"),
        ("Data-efficient carbon capture", "Used active learning to explore open-copper-site MOFs for precombustion CO2 capture with fewer expensive simulations.", "https://doi.org/10.1016/j.cej.2025.167021"),
        ("Transferable materials AI", "Developed cross-diversity and stability-aware models designed to perform beyond a single training database.", "https://doi.org/10.1021/acs.jctc.4c01478"),
        ("Chemically interpretable descriptors", "Created Pore+ descriptors connecting pore geometry and chemistry for trace iodide capture.", "https://doi.org/10.1016/j.seppur.2024.130933"),
    ]
    for title, detail, link in projects:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        set_font(p.add_run(title + ": "), 9.1, bold=True)
        set_font(p.add_run(detail + " "), 9.1)
        add_hyperlink(p, "publication", link, size=9.1)
    heading(doc, "Education and Credentials")
    role_entry(doc, "2023", "PhD in Chemistry", "University of Liverpool, United Kingdom", "Supervisors: Prof. Andrew I. Cooper and Prof. Lifeng Ding.")
    role_entry(doc, "2018", "BSc in Chemistry", "Nanjing Tech University, China")
    body(doc, "Member of the Royal Society of Chemistry (MRSC).", 9.0)
    heading(doc, "Selected Publications")
    selected_dois = {
        "https://doi.org/10.1002/adfm.202519565",
        "https://doi.org/10.1016/j.cej.2025.167021",
        "https://doi.org/10.1021/acs.jctc.4c01478",
        "https://doi.org/10.1039/d4sc05616g",
        "https://doi.org/10.1039/D5TA01139F",
    }
    selected = [p for p in publications if p["doi_or_link"] in selected_dois]
    pub_num_id = add_numbering(doc, "decimal", "%1.", 340, 260)
    for index, paper in enumerate(selected, 1):
        add_publication(doc, paper, index)
    heading(doc, "Additional Leadership")
    for item in [
        "Mentored five student researchers across doctoral, master's, and undergraduate projects, including work that led to peer-reviewed publication.",
        "Contributed to multi-institution research programmes and secured large-scale national and university CPU/GPU allocations.",
        "Frequent peer reviewer for leading journals in chemical engineering, materials, computation, and environmental science.",
    ]:
        bullet(doc, item, bullet_id, 9.0)
    return doc


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    publications = read_publications()
    academic = build_academic(publications)
    industry = build_industry(publications)
    academic.core_properties.title = "Xiaoyu Wu - Academic Curriculum Vitae"
    industry.core_properties.title = "Xiaoyu Wu - Industry Resume"
    academic.core_properties.author = "Xiaoyu Wu"
    industry.core_properties.author = "Xiaoyu Wu"
    academic.save(OUT / "Xiaoyu_Wu_Academic_CV.docx")
    industry.save(OUT / "Xiaoyu_Wu_Industry_CV.docx")
    print(OUT)


if __name__ == "__main__":
    main()
